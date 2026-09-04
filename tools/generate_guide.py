#!/usr/bin/env python3
"""
Generate Abu Twins ATOMS System 101 Training & Operations Guide (.docx)
Designed with high visual aesthetics, 5th-grade friendly language,
Ibadan branch context, and complete feature coverage.
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# --- Color Palette Constants ---
NAVY_HEX = "1B365D"      # Abu Twins Deep Royal Navy (Primary Header)
GOLD_HEX = "C59B27"      # Elegant Abu Twins Gold Accent
CHARCOAL_HEX = "222222"  # Crisp reading text
LIGHT_BG_HEX = "F4F7FB"  # Soft blue/grey table header & callout bg
BORDER_HEX = "D1D5DB"    # Clean subtle table border
SUCCESS_BG = "E6F4EA"    # Soft green
SUCCESS_TEXT = "137333"
WARN_BG = "FEF7E0"       # Soft yellow
WARN_TEXT = "B06000"
ALERT_BG = "FCE8E6"      # Soft red/pink
ALERT_TEXT = "C5221F"
BLUE_BG = "E8F0FE"       # Soft informational blue
BLUE_TEXT = "1A73E8"

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table, border_color="D1D5DB"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def add_callout(doc, title, text, box_type="info"):
    """Adds a stylish shaded callout card with an icon and border."""
    bg_color = BLUE_BG
    text_color = BLUE_TEXT
    icon = "ℹ️"
    
    if box_type == "example":
        bg_color = LIGHT_BG_HEX
        text_color = NAVY_HEX
        icon = "📍 REAL LIFE AT ABU TWINS:"
    elif box_type == "warning":
        bg_color = WARN_BG
        text_color = WARN_TEXT
        icon = "⚠️ IMPORTANT RULE:"
    elif box_type == "tip":
        bg_color = SUCCESS_BG
        text_color = SUCCESS_TEXT
        icon = "💡 PRO TIP:"
    elif box_type == "golden":
        bg_color = "FFF8E1"
        text_color = "8D6E63"
        icon = "🌟 GOLDEN RULE:"

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
    
    # Left border in accent color
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{text_color}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    r_title = p.add_run(f"{icon} {title}\n")
    r_title.bold = True
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(11)
    r_title.font.color.rgb = RGBColor.from_string(text_color)
    
    r_text = p.add_run(text)
    r_text.font.name = "Calibri"
    r_text.font.size = Pt(10.5)
    r_text.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def format_heading_1(doc, text):
    h = doc.add_heading(level=1)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    r = h.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY_HEX)
    return h

def format_heading_2(doc, text):
    h = doc.add_heading(level=2)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True
    r = h.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(GOLD_HEX)
    return h

def format_heading_3(doc, text):
    h = doc.add_heading(level=3)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(3)
    h.paragraph_format.keep_with_next = True
    r = h.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    return h

def add_clean_paragraph(doc, text, bold_prefix=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix + " ")
        r_pre.bold = True
        r_pre.font.name = "Calibri"
        r_pre.font.size = Pt(11)
        r_pre.font.color.rgb = RGBColor.from_string(NAVY_HEX)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return p

def add_bullet_point(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix + " ")
        r_pre.bold = True
        r_pre.font.name = "Calibri"
        r_pre.font.size = Pt(10.5)
        r_pre.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return p

def build_document():
    doc = docx.Document()
    
    # Page setup - 0.8 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # --- COVER / TITLE BLOCK ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(4)
    
    r_sub0 = title_p.add_run("ABU TWINS ENTERPRISES\n")
    r_sub0.font.name = "Calibri"
    r_sub0.font.size = Pt(14)
    r_sub0.bold = True
    r_sub0.font.color.rgb = RGBColor.from_string(GOLD_HEX)
    
    r_title = title_p.add_run("ATOMS SYSTEM 101\n")
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(28)
    r_title.bold = True
    r_title.font.color.rgb = RGBColor.from_string(NAVY_HEX)
    
    r_sub = title_p.add_run("The Complete Staff Operations & Workflow Guide\n")
    r_sub.font.name = "Calibri"
    r_sub.font.size = Pt(16)
    r_sub.bold = True
    r_sub.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    
    r_sub2 = title_p.add_run("Simple, Clear & Step-by-Step Training Manual for All Branches Across Ibadan\n(Iwo Road, Challenge, Dugbe, Bodija, Mokola & Central Vault)")
    r_sub2.font.name = "Calibri"
    r_sub2.font.size = Pt(11)
    r_sub2.italic = True
    r_sub2.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    add_callout(doc, "WELCOME TO ABU TWINS ATOMS!", 
        "This guide is written in very simple everyday language to help every member of our team — from our new sales officers at Challenge to our vault managers at Iwo Road — master our system effortlessly. You don't need any special computer degree! Just read along and follow the clear steps.",
        box_type="tip")
        
    doc.add_page_break()

    # --- TABLE OF CONTENTS SUMMARY ---
    format_heading_1(doc, "Table of Contents & Quick Navigation")
    
    toc_items = [
        ("Chapter 1", "What is ATOMS and Why Do We Use It?", "The 3 Golden Rules of Abu Twins"),
        ("Chapter 2", "User Roles & Permissions", "Who Does What & Who Sees What"),
        ("Chapter 3", "System Basics & Getting Started", "Logging in, Mobile App, Branch Selection & Dashboard"),
        ("Chapter 4", "Stock & Inventory Management", "Adding Phones, Receiving Shipments, IMEIs & Audits"),
        ("Chapter 5", "Branch Transfers Workflow", "Moving Phones Safely Between Iwo Road, Challenge, etc."),
        ("Chapter 6", "Sales & Point of Sale (POS)", "Scanning IMEIs, Payment Types, Receipts & WhatsApp"),
        ("Chapter 7", "The Minimum Price Protection & Approvals", "Why Discounts Need Manager Sign-Off"),
        ("Chapter 8", "Phone Swaps & Trade-Ins", "The Complete Trade-In Calculation & Workflow"),
        ("Chapter 9", "Customer Returns & Warranty Claims", "Automatic Invoice Lookup & Warranty Checks"),
        ("Chapter 10", "Repairs & Workshop Operations", "Engineers, Spare Parts & Tracking Fixes"),
        ("Chapter 11", "Expenses, Cash Drawer & Supplier Ledgers", "Handling Daily Shop Money & Debts"),
        ("Chapter 12", "The Approvals Desk", "How Managers Review & Decide on Requests"),
        ("Chapter 13", "Offline Selling (No Internet? No Problem!)", "How ATOMS Queues & Syncs Transactions"),
        ("Chapter 14", "Reports, Daily Digests & Analytics", "End of Day Numbers & Staff Leaderboards"),
        ("Chapter 15", "Frequently Asked Questions (FAQ) & Floor Scenarios", "Real-Life Problem Solving"),
    ]
    
    tbl_toc = doc.add_table(rows=1, cols=3)
    tbl_toc.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl_toc)
    
    hdr_cells = tbl_toc.rows[0].cells
    hdr_cells[0].text = "Chapter"
    hdr_cells[1].text = "Topic"
    hdr_cells[2].text = "Key Focus"
    for c in hdr_cells:
        set_cell_background(c, LIGHT_BG_HEX)
        set_cell_margins(c)
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(NAVY_HEX)
        
    for ch, topic, focus in toc_items:
        row = tbl_toc.add_row()
        r_cells = row.cells
        r_cells[0].text = ch
        r_cells[1].text = topic
        r_cells[2].text = focus
        r_cells[0].paragraphs[0].runs[0].font.bold = True
        for c in r_cells:
            set_cell_margins(c)
            c.paragraphs[0].runs[0].font.size = Pt(10)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- CHAPTER 1 ---
    format_heading_1(doc, "Chapter 1: What is ATOMS and Why Do We Use It?")
    add_clean_paragraph(doc, "Think of ATOMS as the super-smart brain and memory of Abu Twins. In the past, shops used paper notebooks or simple spreadsheets to write down sales, phone serial numbers, and customer debts. But paper gets lost, numbers get mixed up, and nobody knows for sure how many iPhones are left in the glass display at Challenge or in the big vault at Iwo Road.")
    add_clean_paragraph(doc, "ATOMS solves all of that! It connects all our branches across Ibadan together in real time so that our business runs smoothly, safely, and cleanly.")
    
    format_heading_2(doc, "The 3 Golden Rules of ATOMS")
    add_clean_paragraph(doc, "Every staff member must remember these three simple rules every single day:")
    
    add_bullet_point(doc, "Every single smartphone has a unique 15-digit IMEI number (its digital fingerprint). We NEVER sell, transfer, or accept a phone without scanning or typing its exact IMEI into ATOMS. This ensures no phone ever disappears.", bold_prefix="1. The IMEI is King:")
    add_bullet_point(doc, "In ATOMS, once a sale, payment, or transfer is completed, nobody can secretly delete it or change the numbers. If a genuine mistake happened, we don't 'rub it off' — we make an official correction (called a Void or Reversal) with manager approval. This protects everyone from false blame!", bold_prefix="2. We Never Erase History:")
    add_bullet_point(doc, "Every phone that leaves the shelf must have matching money collected (or an approved debt recorded). Every naira in the cash drawer must match the receipts in the system at closing time.", bold_prefix="3. Stock & Money Must Balance:")
    
    add_callout(doc, "WHY THIS MATTERS TO YOU", 
        "When you use ATOMS properly, your work is 100% transparent. If a customer claims they bought a faulty phone from you 3 months ago, you don't have to argue — you just type the IMEI into ATOMS and see the exact second it was sold, who sold it, and whether the warranty is still valid!",
        box_type="golden")

    # --- CHAPTER 2 ---
    format_heading_1(doc, "Chapter 2: User Roles & Permissions (Who Does What?)")
    add_clean_paragraph(doc, "Just like a football team has a goalkeeper, defenders, midfielders, and strikers who each have a special job, ATOMS gives each staff member the exact tools they need to do their job well without confusing buttons they don't need.")
    
    format_heading_2(doc, "Detailed Role Breakdown")
    
    roles_info = [
        ("1. ATOMS CEO & Director", "Executive Leadership", 
         "Has full access to everything across all branches in Ibadan. Sees company-wide profits, approves major company expenses, reviews branch performance, and manages staff user accounts."),
        ("2. ATOMS Auditor", "Inspection & Quality Control", 
         "Checks stock balances across all branches. Has special power to review and approve Stock Count Adjustments when physical shelf counts are verified against system records. Cannot sell or discount phones."),
        ("3. ATOMS Accountant", "Financial & Ledger Master", 
         "Monitors bank transfers, customer debts, daily cash totals across all shops, and payments to phone suppliers. Approves shop expenses above ₦50,000."),
        ("4. ATOMS Branch Manager", "Branch Floor Leader (Iwo Road, Challenge, etc.)", 
         "Runs daily branch operations. Approves customer price discounts below minimum floor price, approves staff expenses, initiates stock transfers to other branches, and signs off on daily branch Z-reports."),
        ("5. ATOMS Vault Manager", "Central Warehouse Custodian", 
         "Guards the central stock repository (e.g. at Iwo Road Central Vault). Receives new shipments of phones arriving from Lagos or overseas, registers IMEIs into the system, and supplies phones to other branches when requested."),
        ("6. ATOMS Inbound Coordinator", "Shipment Receiver & Inspector", 
         "Assists in unboxing new carton deliveries, cross-checking physical boxes against supplier packing lists, and entering incoming manifests into the system."),
        ("7. ATOMS Cashier", "Money & Payment Officer", 
         "Receives cash, verifies bank transfer alerts, swipes POS debit cards, splits payments, issues official printed receipts, and balances the physical cash drawer at the end of each shift."),
        ("8. ATOMS Sales Officer", "Customer Relationship & POS Champion", 
         "Welcomes customers on the sales floor, shows phone models, scans IMEIs to start sales, evaluates trade-in phones for swaps, and creates customer invoices."),
        ("9. ATOMS Engineer", "Technical Repair Specialist", 
         "Manages the workshop (e.g. at Mokola / Challenge). Receives faulty phones, diagnoses screen/battery issues, logs spare parts costs, updates repair progress, and hands fixed phones back to customers."),
        ("10. ATOMS Inventory Officer", "Stock Checker & Shelf Arranger", 
         "Performs daily morning stock counts, verifies shelf display IMEIs, and reports any damaged or misplaced boxes to the Branch Manager."),
    ]
    
    tbl_roles = doc.add_table(rows=1, cols=3)
    tbl_roles.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl_roles)
    
    r_hdr = tbl_roles.rows[0].cells
    r_hdr[0].text = "Role Title"
    r_hdr[1].text = "Main Focus"
    r_hdr[2].text = "What They Can Do & See"
    for c in r_hdr:
        set_cell_background(c, LIGHT_BG_HEX)
        set_cell_margins(c)
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(NAVY_HEX)
        
    for title, focus, desc in roles_info:
        row = tbl_roles.add_row()
        c_list = row.cells
        c_list[0].text = title
        c_list[1].text = focus
        c_list[2].text = desc
        c_list[0].paragraphs[0].runs[0].font.bold = True
        for c in c_list:
            set_cell_margins(c)
            c.paragraphs[0].runs[0].font.size = Pt(9.5)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    add_callout(doc, "BRANCH ACCESS IS SAFEGUARDED", 
        "If you are assigned as a Sales Officer or Branch Manager at the Challenge Branch, you will only see and sell phones physically present at the Challenge branch! You cannot accidentally sell a phone that is sitting in the glass case at Iwo Road or Dugbe.",
        box_type="warning")

    # --- CHAPTER 3 ---
    format_heading_1(doc, "Chapter 3: System Basics & Getting Started")
    add_clean_paragraph(doc, "Getting started with ATOMS is as easy as opening WhatsApp or browsing the web. You can use ATOMS on a desktop computer, a tablet, or right on your smartphone!")
    
    format_heading_2(doc, "1. How to Log In")
    add_bullet_point(doc, "Open your web browser (Chrome or Safari) and go to your shop's ATOMS link (e.g. on your shop computer or phone).", bold_prefix="Step 1:")
    add_bullet_point(doc, "Enter your assigned Username and Password. Never share your password with anyone else — every action you take is stamped with your name.", bold_prefix="Step 2:")
    add_bullet_point(doc, "If you have access to multiple branches, pick the branch you are working in today (e.g. 'Iwo Road Branch' or 'Challenge Branch').", bold_prefix="Step 3:")
    
    format_heading_2(doc, "2. Installing the ATOMS Mobile App on Your Phone")
    add_clean_paragraph(doc, "ATOMS has a special Progressive Web App (PWA) built specifically for your smartphone:")
    add_bullet_point(doc, "Open the link on your mobile phone browser.", bold_prefix="On Android:")
    add_bullet_point(doc, "Tap the 3 dots menu at the top right and select 'Add to Home Screen' or 'Install App'.", bold_prefix="On iPhone:")
    add_bullet_point(doc, "Tap the Share button (the square with an arrow pointing up) and tap 'Add to Home Screen'.", bold_prefix="Result:")
    add_bullet_point(doc, "You now have the official Abu Twins ATOMS icon on your phone home screen just like an app from the App Store!", bold_prefix="Done:")

    format_heading_2(doc, "3. Understanding the Home Dashboard")
    add_clean_paragraph(doc, "When you log in, your home dashboard greets you with simple cards that summarize the day:")
    add_bullet_point(doc, "How much money this branch has made so far today.", bold_prefix="💰 Today's Revenue:")
    add_bullet_point(doc, "Physical cash currently sitting in the cashier drawer.", bold_prefix="💵 Cash in Drawer:")
    add_bullet_point(doc, "Total number of phones ready on the shelf for customers.", bold_prefix="📦 Available Stock:")
    add_bullet_point(doc, "Any discount requests or expenses waiting for the manager's green light.", bold_prefix="⏳ Pending Approvals:")
    add_bullet_point(doc, "Phone models that are running out (e.g. only 1 iPhone 13 128GB left!).", bold_prefix="⚠️ Low Stock Warnings:")

    format_heading_2(doc, "4. The Universal Search Bar (Find Anything in 1 Second!)")
    add_clean_paragraph(doc, "At the top of every screen is a universal search bar. You can type:")
    add_bullet_point(doc, "Shows the phone's full life story — when it arrived, what supplier brought it, what branch it moved to, what invoice sold it, and remaining warranty.", bold_prefix="Any 15-Digit IMEI:")
    add_bullet_point(doc, "Pulls up the receipt, payment breakdown, and allows instant reprinting or WhatsApp sending.", bold_prefix="An Invoice Number (e.g. INV-2026-0042):")
    add_bullet_point(doc, "Shows all past purchases, total money spent, and any pending debt balance.", bold_prefix="A Customer's Phone Number or Name:")

    # --- CHAPTER 4 ---
    format_heading_1(doc, "Chapter 4: Stock & Inventory Management (From Shipment to Shelf)")
    add_clean_paragraph(doc, "Phones are the heart of Abu Twins. Here is how phones enter our business, get tracked, and stay secure.")
    
    format_heading_2(doc, "1. Adding Products & Storage Variants")
    add_clean_paragraph(doc, "Before phones arrive, we define the product in ATOMS. A single phone model can have different colors and storage sizes:")
    add_bullet_point(doc, "e.g. 'Apple iPhone 15 Pro Max'", bold_prefix="Model Name:")
    add_bullet_point(doc, "e.g. 256GB, 512GB, 1TB.", bold_prefix="Storage Options:")
    add_bullet_point(doc, "e.g. Natural Titanium, Blue Titanium, Black Titanium.", bold_prefix="Color Options:")
    add_bullet_point(doc, "The absolute lowest price staff can sell this phone for without manager approval (e.g. ₦1,450,000).", bold_prefix="Minimum Floor Price:")
    add_bullet_point(doc, "Standard is 365 days (1 year) for new devices, or 90 days for UK-used devices.", bold_prefix="Warranty Days:")
    add_bullet_point(doc, "e.g. Alert me when fewer than 2 units remain in this branch.", bold_prefix="Low Stock Threshold:")

    format_heading_2(doc, "2. Receiving Shipments (Goods-In / Inbound Manifest)")
    add_clean_paragraph(doc, "When a new carton of phones arrives from Lagos or an overseas supplier at the Central Vault (Iwo Road):")
    add_bullet_point(doc, "The Inbound Coordinator opens 'Goods-In' and selects the Supplier (e.g. 'Emeka Lagos Gadgets' or 'Shenzhen Tech').", bold_prefix="Step 1:")
    add_bullet_point(doc, "Selects the Product Model, Storage, and Color.", bold_prefix="Step 2:")
    add_bullet_point(doc, "Scans the 15-digit IMEI on each phone box using a barcode scanner or phone camera. (You can also type it manually).", bold_prefix="Step 3:")
    add_bullet_point(doc, "Enter the cost price per unit. ATOMS automatically adds the devices to available stock and updates the supplier's balance in our accounting ledger!", bold_prefix="Step 4:")

    add_callout(doc, "DUPLICATE PROTECTION", 
        "If you accidentally try to scan the same IMEI twice, ATOMS will immediately beep and warn: 'This IMEI already exists in the system!' You can never double-count stock by mistake.",
        box_type="warning")

    format_heading_2(doc, "3. Stock Counts & Audits (The Weekly Shelf Check)")
    add_clean_paragraph(doc, "To make sure no phone went missing or was misplaced, the Inventory Officer and Branch Manager do a periodic Stock Count:")
    add_bullet_point(doc, "The manager clicks 'Start Stock Count' in ATOMS for their branch.", bold_prefix="Step 1:")
    add_bullet_point(doc, "The system takes a silent snapshot of what is supposed to be on the shelf.", bold_prefix="Step 2:")
    add_bullet_point(doc, "The officer walks along the glass cases and scans every physical phone.", bold_prefix="Step 3:")
    add_bullet_point(doc, "ATOMS compares the physical scan against the system records. If all match (Variance = 0), the count is closed with a green checkmark! If any phone is missing or extra, an alert is sent to the Auditor for investigation.", bold_prefix="Step 4:")

    # --- CHAPTER 5 ---
    format_heading_1(doc, "Chapter 5: Branch Transfers Workflow (Moving Phones Safely)")
    add_clean_paragraph(doc, "Abu Twins has multiple branches across Ibadan. Customers often want a specific color or storage that is currently at another branch. Here is how we transfer phones safely:")

    add_callout(doc, "REAL EXAMPLE: IWO ROAD TO CHALLENGE TRANSFER", 
        "A customer is standing at the Challenge Branch asking for an iPhone 14 Pro Max 512GB Deep Purple. Challenge has zero in stock, but Iwo Road Central Vault has 3 pieces. The Challenge Manager calls Iwo Road, and a transfer is created.",
        box_type="example")

    format_heading_2(doc, "The 3 Steps of a Safe Transfer")
    add_bullet_point(doc, "The sending branch (Iwo Road) opens 'Stock Transfers' → 'New Transfer'. Selects Destination: 'Challenge Branch'. Scans the exact IMEI of the phone leaving the shelf.", bold_prefix="1. Initiating the Transfer:")
    add_bullet_point(doc, "The status changes to 'In Transit'. The phone is no longer available to be sold at Iwo Road, but it is not yet available at Challenge. It is marked as being with the dispatch rider.", bold_prefix="2. On The Road (In-Transit):")
    add_bullet_point(doc, "When the dispatch rider arrives at Challenge, the Challenge Manager opens ATOMS, scans the phone's IMEI to verify it is the exact phone sent, and clicks 'Receive'. The phone instantly becomes available in Challenge stock!", bold_prefix="3. Receiving at Destination:")

    add_callout(doc, "WHAT IF A PHONE IS DAMAGED ON THE ROAD?", 
        "If a phone arrives with a cracked screen or the rider brought the wrong device, the receiving manager clicks 'Report Issue' instead of 'Receive'. This immediately notifies the Auditor and CEO with a complete time-stamped record.",
        box_type="warning")

    # --- CHAPTER 6 ---
    format_heading_1(doc, "Chapter 6: Sales & Point of Sale (POS)")
    add_clean_paragraph(doc, "The Point of Sale (POS) screen is where Sales Officers and Cashiers spend most of their day making customers happy.")

    format_heading_2(doc, "The 4 Easy Steps to Complete a Sale")
    add_bullet_point(doc, "Tap 'New Sale'. Point the camera or barcode scanner at the phone's box or screen to scan the 15-digit IMEI. The system automatically identifies the exact model, color, storage, and standard retail price.", bold_prefix="Step 1: Scan IMEI:")
    add_bullet_point(doc, "Type the customer's phone number. If they are an existing customer, their name and history pop up automatically. If they are new, type their Name and WhatsApp number.", bold_prefix="Step 2: Customer Info:")
    add_bullet_point(doc, "Standard sales are 'Retail'. If selling in bulk to a partner dealer in Ibadan, choose 'Wholesale'.", bold_prefix="Step 3: Sale Type:")
    add_bullet_point(doc, "Enter how the customer is paying: Cash, Bank Transfer, POS/Card, or Split Payment.", bold_prefix="Step 4: Collect Payment & Print:")

    format_heading_2(doc, "Supported Payment Methods")
    add_bullet_point(doc, "Cash counted and verified by the cashier.", bold_prefix="💵 Cash:")
    add_bullet_point(doc, "Direct transfer to Abu Twins bank account. Always verify bank alert before handing over device!", bold_prefix="🏦 Bank Transfer:")
    add_bullet_point(doc, "Customer swipes their debit/credit card on the shop POS terminal.", bold_prefix="💳 POS / Card:")
    add_bullet_point(doc, "e.g. A customer buying a ₦800,000 phone pays ₦500,000 by bank transfer and ₦300,000 in cash. ATOMS handles this smoothly with one click!", bold_prefix="🔀 Split Payment:")
    add_bullet_point(doc, "If an approved corporate client or VIP is paying part now and part next week, ATOMS records the paid amount and automatically tracks the remaining debt balance.", bold_prefix="📝 Part Payment / Customer Debt:")

    format_heading_2(doc, "Receipts & WhatsApp Integration")
    add_clean_paragraph(doc, "Once the sale is posted, ATOMS gives you two instant options:")
    add_bullet_point(doc, "Click 'Print Receipt' for a clean 80mm thermal receipt featuring the official Abu Twins logo, branch address, IMEI, price, and warranty terms.", bold_prefix="1. Thermal Print:")
    add_bullet_point(doc, "Click 'Send via WhatsApp' to automatically open a friendly chat with the customer containing their digital invoice and warranty certificate!", bold_prefix="2. 1-Click WhatsApp Send:")

    # --- CHAPTER 7 ---
    format_heading_1(doc, "Chapter 7: Minimum Price Protection & Approvals")
    add_clean_paragraph(doc, "In Nigeria, bargaining is part of shopping! Customers will often say, 'Oga, give me last price!' But how do we prevent staff from selling at a loss?")

    format_heading_2(doc, "How the Price Guard Works")
    add_clean_paragraph(doc, "Every phone model in ATOMS has an official Floor Price (Minimum Price) set by the CEO. For example:")
    add_bullet_point(doc, "₦750,000", bold_prefix="Standard Selling Price:")
    add_bullet_point(doc, "₦720,000", bold_prefix="Minimum Allowed Floor Price:")
    
    add_clean_paragraph(doc, "If a Sales Officer enters a price of ₦730,000, the sale goes through immediately because it is above the ₦720,000 floor.")
    add_clean_paragraph(doc, "However, if the customer insists on paying ₦700,000 (which is ₦20,000 below the floor):")
    add_bullet_point(doc, "The POS does NOT crash or say 'Error'.", bold_prefix="What Happens:")
    add_bullet_point(doc, "The system creates a 'Pending Approval' request.", bold_prefix="Manager Alert:")
    add_bullet_point(doc, "The phone is temporarily reserved for that customer.", bold_prefix="Stock Reserved:")
    add_bullet_point(doc, "The Branch Manager or CEO sees the alert on their dashboard: 'Sales Officer Kunle at Challenge is asking to sell iPhone 13 for ₦700,000 (Floor: ₦720,000)'.", bold_prefix="Manager Action:")
    add_bullet_point(doc, "If the manager clicks 'Approve', the sale completes instantly and the receipt prints!", bold_prefix="Result:")

    # --- CHAPTER 8 ---
    format_heading_1(doc, "Chapter 8: Phone Swaps & Trade-Ins (Abu Twins Speciality!)")
    add_clean_paragraph(doc, "Abu Twins is famous across Ibadan for smooth phone upgrades and swaps. ATOMS has a dedicated Trade-In Calculator that makes this completely foolproof.")

    add_callout(doc, "REAL EXAMPLE: TRADE-IN AT BODIJA BRANCH", 
        "Customer Tunde walks into our Bodija branch holding his used iPhone 11 (64GB Black). He wants to upgrade to a brand new iPhone 13 (128GB Pink) selling for ₦650,000.",
        box_type="example")

    format_heading_2(doc, "Step-by-Step Swap Process")
    add_bullet_point(doc, "Open 'Trade-In / Swap Desk'. Enter customer name and phone number.", bold_prefix="Step 1:")
    add_bullet_point(doc, "Inspect Tunde's iPhone 11 (screen, battery health, Face ID, body). Agree on its trade-in valuation (e.g. ₦200,000). Scan or type its IMEI. ATOMS registers this used phone directly into Abu Twins inventory!", bold_prefix="Step 2: The Trade-In Phone:")
    add_bullet_point(doc, "Scan the IMEI of the brand new iPhone 13 from the shelf (₦650,000).", bold_prefix="Step 3: The New Phone:")
    add_bullet_point(doc, "ATOMS automatically calculates: ₦650,000 (New Phone) - ₦200,000 (Trade-In Allowance) = ₦450,000 Top-Up Difference.", bold_prefix="Step 4: The Math:")
    add_bullet_point(doc, "Tunde pays the ₦450,000 difference by bank transfer or cash. The receipt prints showing the full swap breakdown, and both inventory records update in one click!", bold_prefix="Step 5: Payment & Receipt:")

    # --- CHAPTER 9 ---
    format_heading_1(doc, "Chapter 9: Customer Returns & Warranty Claims")
    add_clean_paragraph(doc, "Every phone sold at Abu Twins comes with peace of mind. But what happens when a customer returns with a complaint?")

    format_heading_2(doc, "1. The 1-Scan Lookup")
    add_clean_paragraph(doc, "You never have to search through old paper books. Just open 'Customer Returns' and scan the phone's IMEI:")
    add_bullet_point(doc, "The exact original invoice number.", bold_prefix="ATOMS Instantly Displays:")
    add_bullet_point(doc, "The date and time it was purchased.", bold_prefix="Sale Date:")
    add_bullet_point(doc, "Which staff member sold it and at what branch.", bold_prefix="Sold By:")
    add_bullet_point(doc, "How many warranty days are still remaining.", bold_prefix="Warranty Status:")

    format_heading_2(doc, "2. Choosing the Return Action")
    add_bullet_point(doc, "If the phone is within warranty and has a hardware fault, ATOMS routes the phone directly to our Engineer Workshop and opens a Repair Ticket.", bold_prefix="Option A: Send for Warranty Repair:")
    add_bullet_point(doc, "Swap the customer to a working replacement device from stock.", bold_prefix="Option B: Direct Replacement:")
    add_bullet_point(doc, "If approved by management, the customer ledger is credited or cash refunded, and the returned phone is returned to stock or marked as faulty.", bold_prefix="Option C: Refund / Return:")

    # --- CHAPTER 10 ---
    format_heading_1(doc, "Chapter 10: Repairs & Workshop Operations")
    add_clean_paragraph(doc, "Our Engineer Department (e.g. at Mokola / Challenge) handles repairs for both customer phones and internal branch stock.")

    format_heading_2(doc, "The Repair Ticket Lifecycle")
    add_bullet_point(doc, "The engineer or sales officer enters the phone IMEI/Serial, customer name, and customer complaint (e.g. 'Broken screen, battery draining fast'). Status: 'Received'.", bold_prefix="1. Ticket Creation:")
    add_bullet_point(doc, "The engineer tests the device, lists required parts (e.g. 'Original OLED Screen + Battery') and repair cost. Status: 'Diagnosing' → 'In Progress'.", bold_prefix="2. Diagnosis & Parts:")
    add_bullet_point(doc, "Once repaired, the engineer runs quality tests (camera, sound, touch, charging). Status: 'Testing' → 'Ready for Pickup'.", bold_prefix="3. Quality Check:")
    add_bullet_point(doc, "Customer arrives, inspects phone, pays repair bill, and collects phone. The ticket is marked 'Completed' and posted to daily revenue.", bold_prefix="4. Pickup & Payment:")

    # --- CHAPTER 11 ---
    format_heading_1(doc, "Chapter 11: Expenses, Cash Drawer & Supplier Ledgers")
    add_clean_paragraph(doc, "Managing money accurately ensures Abu Twins stays profitable and healthy.")

    format_heading_2(doc, "1. Daily Cash Drawer Reconciliation")
    add_clean_paragraph(doc, "At the end of each working day at Iwo Road, Challenge, Dugbe, etc., the Cashier and Branch Manager perform a Cash Balancing:")
    add_bullet_point(doc, "Cash counted by hand in the drawer must equal: Starting Morning Cash + Cash Sales - Approved Cash Expenses.", bold_prefix="The Formula:")
    add_bullet_point(doc, "The manager signs off on the daily Z-Report in ATOMS. If there is any shortage or excess, an explanation note is required.", bold_prefix="Sign-off:")

    format_heading_2(doc, "2. Recording Daily Shop Expenses")
    add_clean_paragraph(doc, "Shops have daily running costs (e.g. diesel/fuel for generator, electricity tokens, cleaning supplies, staff lunch):")
    add_bullet_point(doc, "Click 'Record Expense' → Select Category (e.g. 'Fuel/Generator', 'Office Supplies', 'Logistics').", bold_prefix="How to Record:")
    add_bullet_point(doc, "Enter the amount and attach a short description.", bold_prefix="Description:")
    add_bullet_point(doc, "Small daily expenses up to ₦50,000 post immediately.", bold_prefix="Under ₦50,000:")
    add_bullet_point(doc, "Any major expense over ₦50,000 (e.g. generator servicing or shop repair) goes to the Accountant / CEO for approval first before cash can be taken!", bold_prefix="Over ₦50,000 (The Safety Rule):")

    format_heading_2(doc, "3. Supplier Ledgers (What We Owe Suppliers)")
    add_clean_paragraph(doc, "When we buy 100 iPhones from our supplier on credit or partial deposit, ATOMS maintains a real-time Supplier Ledger:")
    add_bullet_point(doc, "Shows the total value of phones received.", bold_prefix="Total Invoiced:")
    add_bullet_point(doc, "Shows bank transfers sent to the supplier.", bold_prefix="Total Paid:")
    add_bullet_point(doc, "Shows the exact outstanding balance Abu Twins owes at any moment.", bold_prefix="Current Balance:")

    # --- CHAPTER 12 ---
    format_heading_1(doc, "Chapter 12: The Approvals Desk (Safety First)")
    add_clean_paragraph(doc, "The Approvals Desk is where Branch Managers, Auditors, and the CEO review items that require special permission. This keeps our business safe and prevents mistakes.")

    format_heading_2(doc, "What Appears on the Approvals Desk?")
    
    approvals_list = [
        ("1. Below-Minimum Price Discount", "Sales Officer", "Branch Manager or CEO", 
         "Shows the phone model, IMEI, standard price, minimum floor price, requested selling price, and discount amount. Manager can Approve or Decline with a note."),
        ("2. Large Expense (> ₦50,000)", "Cashier / Manager", "Accountant or CEO", 
         "Shows the branch name, expense category, purpose, and amount. Ensures company funds are spent wisely."),
        ("3. Stock Count Adjustments", "Inventory Officer", "Auditor or CEO", 
         "Shows any difference between physical shelf count and system count. Ensures missing or found devices are thoroughly investigated before adjusting the books."),
    ]
    
    tbl_app = doc.add_table(rows=1, cols=4)
    tbl_app.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl_app)
    
    a_hdr = tbl_app.rows[0].cells
    a_hdr[0].text = "Approval Item"
    a_hdr[1].text = "Requested By"
    a_hdr[2].text = "Decided By"
    a_hdr[3].text = "Details Shown on Screen"
    for c in a_hdr:
        set_cell_background(c, LIGHT_BG_HEX)
        set_cell_margins(c)
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(NAVY_HEX)
        
    for item, req, dec, det in approvals_list:
        row = tbl_app.add_row()
        c_list = row.cells
        c_list[0].text = item
        c_list[1].text = req
        c_list[2].text = dec
        c_list[3].text = det
        c_list[0].paragraphs[0].runs[0].font.bold = True
        for c in c_list:
            set_cell_margins(c)
            c.paragraphs[0].runs[0].font.size = Pt(9.5)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- CHAPTER 13 ---
    format_heading_1(doc, "Chapter 13: Offline Selling (No Internet? No Problem!)")
    add_clean_paragraph(doc, "In Ibadan, internet connections can sometimes be slow or temporarily disconnect. With ATOMS, your sales floor NEVER stops!")

    format_heading_2(doc, "How Offline Mode Works Automatically")
    add_bullet_point(doc, "If the network drops, ATOMS displays an orange badge at the top: 'Offline Mode Active'.", bold_prefix="1. Automatic Detection:")
    add_bullet_point(doc, "You can still scan IMEIs, select customers, choose payment methods, and complete sales just like normal!", bold_prefix="2. Keep Selling:")
    add_bullet_point(doc, "ATOMS safely stores the completed sale in a secure offline queue right inside your device.", bold_prefix="3. Secure Offline Queue:")
    add_bullet_point(doc, "The very moment your internet connection returns, ATOMS automatically sends all queued sales to the main database.", bold_prefix="4. Automatic Sync:")
    add_bullet_point(doc, "ATOMS uses special duplicate-protection tags. Even if your internet reconnects 5 times, it will NEVER double-post a sale or charge a customer twice!", bold_prefix="5. Zero Double-Posting:")

    # --- CHAPTER 14 ---
    format_heading_1(doc, "Chapter 14: Reports, Daily Digests & Staff Analytics")
    add_clean_paragraph(doc, "ATOMS does all the math so management and staff can see progress clearly.")

    format_heading_2(doc, "1. The Daily WhatsApp Digest")
    add_clean_paragraph(doc, "Every evening after close of business, ATOMS generates a neat summary sent directly to Executive Management:")
    add_bullet_point(doc, "Total revenue made across all Ibadan branches today.", bold_prefix="Total Daily Sales:")
    add_bullet_point(doc, "Sales breakdown by branch (Iwo Road vs Challenge vs Dugbe vs Bodija).", bold_prefix="Branch Comparison:")
    add_bullet_point(doc, "Total phones sold and top-selling model of the day.", bold_prefix="Top Phone Model:")
    add_bullet_point(doc, "Cash collected, bank transfers verified, and customer debts pending.", bold_prefix="Cash Position:")

    format_heading_2(doc, "2. Staff Sales Leaderboards")
    add_clean_paragraph(doc, "ATOMS celebrates our hardworking team! The system tracks sales by staff member:")
    add_bullet_point(doc, "Number of phones sold by each sales officer.", bold_prefix="Invoices Generated:")
    add_bullet_point(doc, "Total revenue generated.", bold_prefix="Total Revenue:")
    add_bullet_point(doc, "Cash collection efficiency.", bold_prefix="Collection Rate:")

    # --- CHAPTER 15 ---
    format_heading_1(doc, "Chapter 15: Frequently Asked Questions (FAQ) & Floor Scenarios")
    
    faqs = [
        ("Q: What if the barcode scanner won't scan the phone box?", 
         "A: Don't worry! You can always type the 15-digit IMEI manually into the search box. ATOMS will find the phone instantly."),
        ("Q: A customer wants to pay half in cash and half by bank transfer. How do I do it?", 
         "A: On the POS checkout screen, select 'Split Payment'. Type the cash amount in the Cash box and the transfer amount in the Bank Transfer box. When the total equals the invoice amount, click Complete!"),
        ("Q: A customer claims they sent a bank transfer, but our shop phone hasn't received the alert yet. Can I give them the phone?", 
         "A: NO! Company policy requires that high-value electronics must never leave the shop until the transfer is verified by the Cashier or Branch Manager."),
        ("Q: I made a mistake on an invoice (e.g. wrong customer name or wrong payment method). Can I delete it?", 
         "A: In ATOMS, we do not delete invoices. Ask your Branch Manager to VOID the incorrect invoice (stating the reason), and then issue the correct invoice. Everything stays clean and transparent."),
        ("Q: A phone is showing 'Available' in ATOMS, but we can't find it in the glass case at Challenge branch. What should I do?", 
         "A: Search the IMEI in ATOMS to check its transfer history. If it was transferred from Iwo Road, check if the transfer was received or is still in-transit. If still missing, notify your Branch Manager immediately to initiate an audit check."),
        ("Q: What if a customer brings a phone for trade-in that is iCloud locked or has a damaged IMEI?", 
         "A: Abu Twins does NOT accept locked or unverified devices for trade-in. The device must be fully unlocked and tested by the Sales Officer or Engineer before entering the swap valuation into ATOMS."),
    ]
    
    for q, a in faqs:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(8)
        p_q.paragraph_format.space_after = Pt(2)
        r_q = p_q.add_run(q)
        r_q.bold = True
        r_q.font.name = "Calibri"
        r_q.font.size = Pt(11)
        r_q.font.color.rgb = RGBColor.from_string(NAVY_HEX)
        
        p_a = doc.add_paragraph()
        p_a.paragraph_format.space_before = Pt(0)
        p_a.paragraph_format.space_after = Pt(6)
        r_a = p_a.add_run(a)
        r_a.font.name = "Calibri"
        r_a.font.size = Pt(10.5)
        r_a.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # --- QUICK REFERENCE CHEAT SHEET ---
    doc.add_page_break()
    format_heading_1(doc, "Abu Twins ATOMS Quick Reference Card")
    
    add_callout(doc, "DAILY OPENING CHECKLIST", 
        "1. Turn on shop computer / tablet / phone.\n"
        "2. Log in with your personal credentials and select your active branch.\n"
        "3. Check 'Low Stock Warnings' and 'Pending Approvals' on your dashboard.\n"
        "4. Confirm morning cash in drawer with Cashier / Branch Manager.\n"
        "5. Keep barcode scanner or phone camera ready for quick IMEI scanning!",
        box_type="tip")
        
    add_callout(doc, "DAILY CLOSING CHECKLIST", 
        "1. Complete or resolve any pending offline sale queues.\n"
        "2. Reconcile physical cash drawer against ATOMS Cash Position.\n"
        "3. Review and sign off on the Branch Daily Z-Report.\n"
        "4. Verify all received transfers have been scanned and accepted into stock.\n"
        "5. Lock all glass displays and central vault securely. Log out of ATOMS.",
        box_type="golden")

    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_footer.paragraph_format.space_before = Pt(24)
    r_foot = p_footer.add_run("© 2026 Abu Twins Enterprises. All rights reserved. System ATOMS v1.8.0")
    r_foot.font.name = "Calibri"
    r_foot.font.size = Pt(9.5)
    r_foot.italic = True
    r_foot.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    output_path = "/Users/kudirat/Desktop/abutwins-invent/Abu_Twins_ATOMS_System_101_Guide.docx"
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    build_document()
